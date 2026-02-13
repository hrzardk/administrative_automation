import pandas as pd
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- KONFIGURASI FILE ---
INPUT_EXCEL = 'D:\!Huawei\Req XLWO\XLWO dash.xlsm' # Ganti dengan nama file Excel Anda
TEMPLATE_WORD = 'D:\!Huawei\Req XLWO\Cisco Swap Template.docx'
SHEET_NAME = 'Body Email' # Sesuaikan dengan nam sheet di gambar Anda
OUTPUT_FOLDER = 'Hasil_Generate'

def set_font_style(run, font_name='Arial', font_size=11, is_bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = is_bold
    r = run._element
    rPr = r.get_or_add_rPr()
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), font_name)
    fonts.set(qn('w:hAnsi'), font_name)
    rPr.append(fonts)

def xpath_replace(element, placeholder, replacement, font_size=11, is_bold=False):
    """
    METODE NUCLEAR: Mencari teks langsung di level XML (Bypass struktur Word)
    Ini bisa menemukan teks di dalam Floating Table atau Text Box.
    """
    # Cari semua elemen teks (<w:t>) di dalam XML scope ini
    # namespace 'w' adalah standar Word
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    text_nodes = element.xpath(f'.//w:t[contains(text(), "{placeholder}")]')
    
    count = 0
    for node in text_nodes:
        if node.text == placeholder:
            # Ganti teks
            node.text = replacement
            
            # Coba akses parent run (<w:r>) untuk formatting
            # Ini agak 'hacky' karena kita memanipulasi XML langsung
            run = node.getparent()
            
            # Kita tidak bisa menggunakan python-docx 'Run' object dengan mudah di sini
            # Jadi kita biarkan format mengikuti aslinya, atau reset manual jika perlu.
            # Biasanya mempertahankan format asli lebih aman untuk tanggal.
            
            print(f"   [SUKSES] XML Replace: {placeholder} -> {replacement}")
            count += 1
            
    return count

def global_replace(doc, placeholder, replacement, font_size, is_bold):
    """
    Kombinasi metode Python-Docx standar + XML Hacking
    """
    replaced_count = 0
    
    # 1. Cara Standar (Untuk Body Utama yang Rapi)
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.clear()
            run = para.add_run(replacement)
            set_font_style(run, 'Arial', font_size, is_bold)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            replaced_count += 1

    # 2. Cara XML (Untuk Header, Footer, Floating Tables, Text Boxes)
    # Scan Body XML
    replaced_count += xpath_replace(doc.element.body, placeholder, replacement)
    
    # Scan Semua Header/Footer di setiap section
    for section in doc.sections:
        # Header
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header: replaced_count += xpath_replace(header._element, placeholder, replacement)
        # Footer
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer: replaced_count += xpath_replace(footer._element, placeholder, replacement)

    if replaced_count == 0:
        print(f"   [WARNING] Placeholder {placeholder} TIDAK DITEMUKAN dimanapun!")

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "", filename)

def reset_table_width(table):
    table.autofit = True
    table.allow_autofit = True
    tblPr = table._element.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'), 'auto'); tblW.set(qn('w:w'), '0')
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:type'), 'auto'); tcW.set(qn('w:w'), '0')

def generate_documents():
    print("Membaca data Excel...")
    try:
        df = pd.read_excel(INPUT_EXCEL, sheet_name=SHEET_NAME, dtype=str)
        df.columns = df.columns.str.strip()
    except Exception as e:
        print(f"Error: {e}")
        return

    col_map = {
        'id': 'Tower ID*', 'host': 'Hostname*', 'ring': 'Ring*', 
        'title': 'Title', 'ip': 'IP*', 'ne': 'NE Type*'
    }

    if col_map['ring'] not in df.columns:
        print(f"Error: Kolom '{col_map['ring']}' tidak ditemukan.")
        return

    if not os.path.exists(OUTPUT_FOLDER): os.makedirs(OUTPUT_FOLDER)

    grouped = df.groupby(col_map['ring'])
    print(f"Ditemukan {len(grouped)} Ring unik.")
    
    # Format Tanggal: "30-Dec" (Sesuai Request)
    today_str = datetime.now().strftime("%d-%b")

    for ring_name, group_data in grouped:
        raw_title = group_data.iloc[0].get(col_map['title'], f"Doc - {ring_name}")
        doc_title = str(raw_title) if pd.notna(raw_title) else f"Doc - {ring_name}"
        
        print(f"Processing: {doc_title}...")
        doc = Document(TEMPLATE_WORD)
        
        # --- 1. REPLACE GLOBAL (XML POWERED) ---
        global_replace(doc, '{{DOC_TITLE}}', doc_title, font_size=25, is_bold=True)
        global_replace(doc, '{{DOC_DATE}}', today_str, font_size=10, is_bold=False)

        # --- 2. ISI TABEL DATA ---
        target_table = None
        for table in doc.tables:
            if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
                if 'Tower ID' in table.rows[0].cells[1].text:
                    target_table = table
                    break
        
        if target_table:
            target_table.style = 'Table Grid'
            for i in range(len(target_table.rows)-1, 0, -1):
                row = target_table.rows[i]
                tbl = row._element.getparent()
                tbl.remove(row._element)

            for idx, row in group_data.iterrows():
                new_row = target_table.add_row()
                row_cells = new_row.cells
                row_num = list(group_data.index).index(idx) + 1
                row_cells[0].text = str(row_num)
                row_cells[1].text = str(row.get(col_map['id'], ''))
                row_cells[2].text = str(row.get(col_map['host'], ''))
                row_cells[3].text = str(row.get(col_map['ring'], ''))
                row_cells[4].text = str(row.get(col_map['ip'], ''))
                row_cells[5].text = str(row.get(col_map['ne'], ''))

                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
            
            reset_table_width(target_table)
        else:
            print(f"Warning: Tabel data utama tidak ditemukan untuk {ring_name}")

        clean_name = sanitize_filename(doc_title)
        output_path = os.path.join(OUTPUT_FOLDER, f"{clean_name}.docx")
        doc.save(output_path)
        print(f"Saved: {output_path}")

    print("\nSelesai! Cek folder:", OUTPUT_FOLDER)

if __name__ == "__main__":
    generate_documents()