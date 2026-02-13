import pandas as pd
import os
import re
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
#      KONFIGURASI PROJECT (EDIT DISINI)
# ==========================================

# GANTI NAMA SCOPE DI SINI SESUAI YANG MAU DI-RUN:
# Pilihan: 'Cisco_Swap', 'Modernize', 'Install_New', 'PAG_Modernize', 'Expansion'
CURRENT_SCOPE = 'PAG_Modernize' 

# DATABASE KONFIGURASI (OTAK AUTOMATION)
PROJECTS = {
    # --- SCOPE 1: CISCO SWAP (Yang sudah jalan) ---
    'Cisco_Swap': {
        'template_file': r'D:\!Huawei\Req XLWO\Code\Cisco Swap\Cisco Swap Template.docx',
        'excel_sheet': 'Body Email',
        'table_keyword': 'Tower ID',   # Kata kunci header tabel Word
        'title_col': 'Title',
        'ring_col': 'Ring*',
        'columns_mapping': [
            'Tower ID*', 'Hostname*', 'Ring*', 'IP*', 'NE Type*'
        ]
    },

    # --- SCOPE 2: MODERNIZE ---
    'Modernize': {
        'template_file': r'D:\!Huawei\Req XLWO\Code\Modernize\Modernize Template.docx',  # Nama file template Anda
        'excel_sheet': 'Sheet1',             # Nama sheet di Excel
        'table_keyword': 'Tower Index A',                  # Header unik di tabel Word Modernize
        'title_col': 'Title',                # Header Excel untuk Judul
        'ring_col': 'Ring',                       # Grouping berdasarkan apa?
        # Urutan kolom Excel yang mau dimasukkan ke Tabel Word (Kiri ke Kanan):
        'columns_mapping': [
            'Ring CSR/MBH', 'Tower Index A', 'Device Name', 'Device IP'
        ]
    },

    # --- SCOPE 3: INSTALL NEW ---
    'Install_New': {
        'template_file': 'template_install.docx',
        'excel_sheet': 'New Installation',
        'table_keyword': 'New Site ID',              # Header unik tabel Install
        'title_col': 'Doc Title',
        'ring_col': 'Region',
        'columns_mapping': [
            'New Site ID', 'Site Name', 'IP Address', 'VLAN ID', 'Eng Name'
        ]
    },

    # --- SCOPE 4: PAG MODERNIZE ---
    'PAG_Modernize': {
        'template_file': r'D:\!Huawei\Req XLWO\Code\PAG Modernize\PAG Modernize.docx',
        'excel_sheet': 'Sheet1',
        'table_keyword': 'Tower Index A',
        'title_col': 'Title',
        'ring_col': 'Ring CSR/MBH',
        'columns_mapping': [
            'Ring CSR/MBH', 'Tower Index A', 'Device Name', 'Device IP'
        ]
    },

    # --- SCOPE 5: EXPANSION ---
    'Expansion': {
        'template_file': r'D:\!Huawei\Req XLWO\Code\Expansion\Template Expansion.docx',
        'excel_sheet': 'Sheet1',
        'table_keyword': 'Tower ID',       # Sesuaikan dengan header tabel Word
        'title_col': 'Title',                   # Kolom untuk judul dokumen
        'ring_col': 'Ring',                     # Kolom untuk grouping
        'region_col': 'Region',                 # Kolom untuk Change Scope (Region)
        'columns_mapping': [                    # Kolom Excel yang diisi ke tabel Word
            'Tower ID', 'PAG Hostname', 'C/AG Tower ID', 'C/AG Hostname', 'Ring', 'Port PAG', 'Port AG'
        ]
    }   
}

INPUT_EXCEL = r'D:\!Huawei\Req XLWO\Code\PAG Modernize\PAG Modernize Template .xlsx' 
# Output folder dengan timestamp (Format: Hasil_Scope_YYYY-MM-DD)
TIMESTAMP = datetime.now().strftime('%Y-%m-%d')
OUTPUT_FOLDER = f'Hasil_{CURRENT_SCOPE.replace(" ", "_")}_{TIMESTAMP}'

# ==========================================
#           CORE ENGINE (JANGAN UBAH)
# ==========================================

def set_font_style(run, font_name='Arial', font_size=11, is_bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = is_bold
    r = run._element
    rPr = r.get_or_add_rPr()
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), font_name)
    fonts.set(qn('w:hAnsi'), font_name)
    rPr.append(fonts)

def xpath_replace(element, placeholder, replacement):
    """Mencari teks di XML level (Floating Tables/Shapes)"""
    # Mencari semua text node yang mengandung placeholder
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    try:
        text_nodes = element.xpath(f'.//w:t', ns)
    except:
        text_nodes = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    
    count = 0
    for node in text_nodes:
        if node.text and placeholder in node.text:
            node.text = node.text.replace(placeholder, replacement)
            count += 1
    return count

def replace_placeholder_in_paragraph(para, placeholder, replacement):
    """Mengganti placeholder yang mungkin terpecah di beberapa run"""
    # Gabungkan semua teks dari runs
    full_text = ''.join([run.text or '' for run in para.runs])
    
    if placeholder not in full_text:
        return False
    
    # Jika placeholder ada di satu run saja
    for run in para.runs:
        if run.text and placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            return True
    
    # Jika placeholder terpecah di beberapa run, gabungkan dan ganti
    new_text = full_text.replace(placeholder, replacement)
    
    # Simpan run pertama dengan teks baru, kosongkan yang lain
    first_run = True
    for run in para.runs:
        if first_run:
            run.text = new_text
            first_run = False
        else:
            run.text = ''
    
    return True

def global_replace(doc, placeholder, replacement, font_size, is_bold):
    """Mengganti placeholder dimanapun (Body, Header, Footer, Table)"""
    replaced_count = 0
    
    # 1. Standard Paragraph (Body Utama)
    for para in doc.paragraphs:
        if placeholder in para.text:
            if replace_placeholder_in_paragraph(para, placeholder, replacement):
                replaced_count += 1
    
    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        if replace_placeholder_in_paragraph(para, placeholder, replacement):
                            replaced_count += 1
    
    # 3. XML Deep Search (Untuk Header, Footer, Floating Tables)
    replaced_count += xpath_replace(doc.element.body, placeholder, replacement)
    
    for section in doc.sections:
        # Header & Footer Check
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header: replaced_count += xpath_replace(header._element, placeholder, replacement)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer: replaced_count += xpath_replace(footer._element, placeholder, replacement)

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "", filename)

def reset_table_width(table):
    """AutoFit to Content Logic"""
    table.autofit = True
    table.allow_autofit = True
    tblPr = table._element.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'), 'auto'); tblW.set(qn('w:w'), '0')
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:type'), 'auto'); tcW.set(qn('w:w'), '0')

def generate_documents():
    print(f"--- MEMULAI AUTOMATION: {CURRENT_SCOPE} ---")
    
    cfg = PROJECTS.get(CURRENT_SCOPE)
    if not cfg:
        print(f"Error: Scope '{CURRENT_SCOPE}' belum disetting di dictionary PROJECTS.")
        return

    print(f"Menggunakan Template: {cfg['template_file']}")
    print("Membaca Excel...")
    try:
        df = pd.read_excel(INPUT_EXCEL, sheet_name=cfg['excel_sheet'], dtype=str)
        df.columns = df.columns.str.strip()
    except Exception as e:
        print(f"Error Membaca Excel/Sheet: {e}")
        return

    # Validasi Kolom
    req_cols = cfg['columns_mapping'] + [cfg['ring_col'], cfg['title_col']]
    missing = [c for c in req_cols if c not in df.columns]
    if missing:
        print(f"Error: Kolom Excel tidak ditemukan: {missing}")
        print(f"Kolom yang ada: {df.columns.tolist()}")
        return

    if not os.path.exists(OUTPUT_FOLDER): os.makedirs(OUTPUT_FOLDER)

    # Strip whitespace dari kolom grouping untuk menghindari duplikasi
    df[cfg['ring_col']] = df[cfg['ring_col']].str.strip()
    
    grouped = df.groupby(cfg['ring_col'])
    print(f"Ditemukan {len(grouped)} Group dokumen.")
    
    today_str = datetime.now().strftime("%d-%b")

    for group_name, group_data in grouped:
        # Ambil Judul
        raw_title = group_data.iloc[0].get(cfg['title_col'], f"Doc - {group_name}")
        doc_title = str(raw_title) if pd.notna(raw_title) else f"Doc - {group_name}"
        
        print(f"Processing: {doc_title}...")
        
        try:
            doc = Document(cfg['template_file'])
        except Exception as e:
            print(f"Error: File template '{cfg['template_file']}' tidak ditemukan.")
            return

        # 1. Replace Judul & Tanggal (Global)
        global_replace(doc, '{{DOC_TITLE}}', doc_title, font_size=25, is_bold=True)
        global_replace(doc, '{{DOC_DATE}}', today_str, font_size=10, is_bold=False)

        # 2. Replace Change Time (Mulai hari ini 23:00 sampai +30 hari 05:00)
        start_date = datetime.now()
        end_date = start_date + timedelta(days=30)
        change_time = f"23:00 {start_date.strftime('%d %b %Y')} - 05:00 {end_date.strftime('%d %b %Y')}"
        global_replace(doc, '{{CHANGE_TIME}}', change_time, font_size=10, is_bold=False)

        # 3. Replace Change Scope (Region)
        if 'region_col' in cfg and cfg['region_col'] in group_data.columns:
            region_val = str(group_data.iloc[0].get(cfg['region_col'], ''))
            if region_val.lower() != 'nan' and region_val:
                change_scope = f"{region_val} region"
            else:
                change_scope = "All region"
        else:
            change_scope = "All region"
        global_replace(doc, '{{CHANGE_SCOPE}}', change_scope, font_size=10, is_bold=False)

        # 2. Isi Tabel Data
        target_table = None
        for table in doc.tables:
            # Cek apakah header tabel mengandung keyword unik scope ini
            if len(table.rows) > 0:
                # Gabungkan semua teks header menjadi satu string untuk pengecekan
                header_row_text = " ".join([cell.text for cell in table.rows[0].cells])
                if cfg['table_keyword'] in header_row_text:
                    target_table = table
                    break
        
        if target_table:
            target_table.style = 'Table Grid'
            
            # Hapus baris kosong (sisakan header)
            for i in range(len(target_table.rows)-1, 0, -1):
                row = target_table.rows[i]
                tbl = row._element.getparent()
                tbl.remove(row._element)

            # Isi Data Baru
            cols_to_fill = cfg['columns_mapping']
            
            for idx, row in group_data.iterrows():
                new_row = target_table.add_row()
                row_cells = new_row.cells
                
                # Handling Kolom "No" (Otomatis)
                # Jika jumlah cell tabel lebih banyak dari jumlah data mapping,
                # Asumsikan kolom pertama adalah Nomor Urut
                start_data_idx = 0
                if len(row_cells) > len(cols_to_fill):
                    row_cells[0].text = str(list(group_data.index).index(idx) + 1)
                    start_data_idx = 1
                
                # Mapping Data Excel ke Kolom Word
                for i, col_name in enumerate(cols_to_fill):
                    if start_data_idx + i < len(row_cells):
                        val = str(row.get(col_name, ''))
                        if val.lower() == 'nan': val = ''
                        row_cells[start_data_idx + i].text = val

                # Formatting Font
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
            
            # 3. Auto Fit
            reset_table_width(target_table)
        else:
            print(f"Warning: Tabel dengan keyword '{cfg['table_keyword']}' tidak ditemukan di template.")

        # Simpan
        clean_name = sanitize_filename(doc_title)
        output_path = os.path.join(OUTPUT_FOLDER, f"{clean_name}.docx")
        doc.save(output_path)
        print(f"Saved: {output_path}")

    print(f"\nSelesai! Hasil ada di folder: {OUTPUT_FOLDER}")

if __name__ == "__main__":
    generate_documents()