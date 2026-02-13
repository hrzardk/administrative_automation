"""
Document Automation Web Interface v2.0
Flask-based web app with Template Manager, Config Editor, Batch Processing, History
"""

from flask import Flask, render_template, request, send_file, jsonify, redirect, url_for
import pandas as pd
import os
import re
import json
import shutil
import zipfile
import uuid
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from werkzeug.utils import secure_filename

app = Flask(__name__)

# ==========================================
#           CONFIG MANAGEMENT
# ==========================================

CONFIG_FILE = 'config.json'
HISTORY_FILE = 'history.json'

def load_config():
    """Load configuration from JSON file"""
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"scopes": {}, "settings": {}}

def save_config(config):
    """Save configuration to JSON file"""
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=4, ensure_ascii=False)

def load_history():
    """Load history from JSON file"""
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"history": []}

def save_history(history):
    """Save history to JSON file"""
    with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
        json.dump(history, f, indent=4, ensure_ascii=False)

def add_to_history(scope, excel_file, doc_count, zip_file, status="success"):
    """Add entry to history"""
    history = load_history()
    entry = {
        "id": str(uuid.uuid4())[:8],
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "scope": scope,
        "excel_file": excel_file,
        "doc_count": doc_count,
        "zip_file": zip_file,
        "status": status
    }
    history["history"].insert(0, entry)
    # Keep only last 50 entries
    history["history"] = history["history"][:50]
    save_history(history)
    return entry

# Load initial config
config = load_config()
settings = config.get("settings", {})

app.config['UPLOAD_FOLDER'] = settings.get('upload_folder', 'uploads')
app.config['OUTPUT_FOLDER'] = settings.get('output_folder', 'output')
app.config['TEMPLATES_FOLDER'] = settings.get('templates_folder', 'templates_store')
app.config['MAX_CONTENT_LENGTH'] = settings.get('max_file_size_mb', 50) * 1024 * 1024

# Ensure folders exist
for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER'], app.config['TEMPLATES_FOLDER']]:
    os.makedirs(folder, exist_ok=True)

# ==========================================
#           HELPER FUNCTIONS
# ==========================================

def xpath_replace(element, placeholder, replacement):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    try:
        text_nodes = element.xpath('.//w:t', ns)
    except:
        text_nodes = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    
    count = 0
    for node in text_nodes:
        if node.text and placeholder in node.text:
            node.text = node.text.replace(placeholder, replacement)
            count += 1
    return count

def replace_placeholder_in_paragraph(para, placeholder, replacement):
    full_text = ''.join([run.text or '' for run in para.runs])
    
    if placeholder not in full_text:
        return False
    
    for run in para.runs:
        if run.text and placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            return True
    
    new_text = full_text.replace(placeholder, replacement)
    first_run = True
    for run in para.runs:
        if first_run:
            run.text = new_text
            first_run = False
        else:
            run.text = ''
    return True

def global_replace(doc, placeholder, replacement):
    replaced_count = 0
    
    for para in doc.paragraphs:
        if placeholder in para.text:
            if replace_placeholder_in_paragraph(para, placeholder, replacement):
                replaced_count += 1
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        if replace_placeholder_in_paragraph(para, placeholder, replacement):
                            replaced_count += 1
    
    replaced_count += xpath_replace(doc.element.body, placeholder, replacement)
    
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header: replaced_count += xpath_replace(header._element, placeholder, replacement)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer: replaced_count += xpath_replace(footer._element, placeholder, replacement)
    
    return replaced_count

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "", filename)

def compress_docx(docx_path, max_size_mb=1):
    """
    Compress DOCX file if it exceeds max_size_mb.
    DOCX is essentially a ZIP archive, so we can recompress it with higher compression.
    
    Args:
        docx_path: Path to the DOCX file
        max_size_mb: Maximum size in MB before compression (default: 1 MB)
    
    Returns:
        tuple: (compressed, original_size, new_size)
    """
    import tempfile
    
    # Check file size
    original_size = os.path.getsize(docx_path)
    max_size_bytes = max_size_mb * 1024 * 1024
    
    if original_size <= max_size_bytes:
        return False, original_size, original_size
    
    try:
        # Create temp directory for extraction
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "temp_compressed.docx")
        
        # Extract the DOCX (which is a ZIP file)
        with zipfile.ZipFile(docx_path, 'r') as zip_in:
            zip_in.extractall(temp_dir)
        
        # Recompress with maximum compression
        with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED, compresslevel=9) as zip_out:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    if file == "temp_compressed.docx":
                        continue
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zip_out.write(file_path, arcname)
        
        # Get new size
        new_size = os.path.getsize(temp_docx)
        
        # Only replace if actually smaller
        if new_size < original_size:
            shutil.copy2(temp_docx, docx_path)
            
        # Cleanup
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        return True, original_size, os.path.getsize(docx_path)
        
    except Exception as e:
        # If compression fails, just return original
        if 'temp_dir' in locals():
            shutil.rmtree(temp_dir, ignore_errors=True)
        return False, original_size, original_size

def reset_table_width(table):
    table.autofit = True
    table.allow_autofit = True
    tblPr = table._element.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None: 
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'auto')
    tblW.set(qn('w:w'), '0')
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None: 
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:type'), 'auto')
            tcW.set(qn('w:w'), '0')

# ==========================================
#           DOCUMENT GENERATION
# ==========================================

def generate_documents(excel_path, scope_key, topology_images=None, progress_callback=None):
    """Generate documents with progress tracking and optional topology images per ring
    
    Args:
        excel_path: Path to Excel file
        scope_key: Scope configuration key
        topology_images: Dict mapping ring names to image paths, e.g. {"Jabo Ring 35": "/path/to/img.png"}
        progress_callback: Optional callback for progress updates
    """
    config = load_config()
    cfg = config.get("scopes", {}).get(scope_key)
    logs = []
    
    # Initialize topology_images as empty dict if None
    if topology_images is None:
        topology_images = {}
    
    def log(msg):
        timestamp = datetime.now().strftime("%H:%M:%S")
        entry = f"[{timestamp}] {msg}"
        logs.append(entry)
        print(entry)
    
    if not cfg:
        log(f"[ERROR] Scope '{scope_key}' tidak ditemukan")
        return None, f"Scope '{scope_key}' tidak ditemukan", logs, 0
    
    log(f"[START] Memulai automation untuk scope: {scope_key}")
    
    # Log topology images info
    if topology_images:
        log(f"[INFO] Topology images: {len(topology_images)} ring(s)")
    
    timestamp = datetime.now().strftime('%Y-%m-%d_%H%M%S')
    output_folder = os.path.join(app.config['OUTPUT_FOLDER'], f'Hasil_{scope_key}_{timestamp}')
    os.makedirs(output_folder, exist_ok=True)
    log(f"[FOLDER] Output: {os.path.basename(output_folder)}")
    
    try:
        df = pd.read_excel(excel_path, sheet_name=cfg['excel_sheet'], dtype=str)
        df.columns = df.columns.str.strip()
        log(f"[EXCEL] Dibaca: {len(df)} baris")
    except Exception as e:
        log(f"[ERROR] Membaca Excel: {str(e)}")
        return None, f"Error membaca Excel: {str(e)}", logs, 0
    
    req_cols = cfg['columns_mapping'] + [cfg['ring_col'], cfg['title_col']]
    missing = [c for c in req_cols if c not in df.columns]
    if missing:
        log(f"[ERROR] Kolom tidak ditemukan: {missing}")
        return None, f"Kolom tidak ditemukan: {missing}", logs, 0
    
    log(f"[OK] Validasi kolom berhasil")
    
    # Normalize Ring column: strip whitespace, normalize multiple spaces, remove leading/trailing spaces
    df[cfg['ring_col']] = df[cfg['ring_col']].fillna('').astype(str).str.strip()
    df[cfg['ring_col']] = df[cfg['ring_col']].str.replace(r'\s+', ' ', regex=True)  # Normalize multiple spaces
    
    # Log unique rings found for debugging
    unique_rings = df[cfg['ring_col']].unique().tolist()
    log(f"[DEBUG] Unique rings found: {unique_rings}")
    
    grouped = df.groupby(cfg['ring_col'])
    total_groups = len(grouped)
    log(f"[INFO] Ditemukan {total_groups} group dokumen")
    
    today_str = datetime.now().strftime("%d-%b")
    generated_files = []
    current = 0
    
    for group_name, group_data in grouped:
        current += 1
        raw_title = group_data.iloc[0].get(cfg['title_col'], f"Doc - {group_name}")
        doc_title = str(raw_title) if pd.notna(raw_title) else f"Doc - {group_name}"
        
        log(f"[{current}/{total_groups}] Processing: {doc_title[:50]}...")
        
        try:
            doc = Document(cfg['template_file'])
        except Exception as e:
            log(f"[ERROR] Template tidak ditemukan: {cfg['template_file']}")
            return None, "Template tidak ditemukan", logs, 0
        
        global_replace(doc, '{{DOC_TITLE}}', doc_title)
        global_replace(doc, '{{DOC_DATE}}', today_str)
        
        start_date = datetime.now()
        end_date = start_date + timedelta(days=30)
        change_time = f"23:00 {start_date.strftime('%d %b %Y')} - 05:00 {end_date.strftime('%d %b %Y')}"
        global_replace(doc, '{{CHANGE_TIME}}', change_time)
        
        # Handle CHANGE_SCOPE placeholder
        if scope_key in ['Change_SFP', 'Upgrade_License'] and 'NE Type*' in group_data.columns:
            # Special handling for Change SFP and Upgrade License: show NE Type count summary
            ne_type_counts = group_data['NE Type*'].value_counts()
            total_devices = ne_type_counts.sum()
            change_scope_lines = [f"Total {total_devices} Device:"]
            for ne_type, count in ne_type_counts.items():
                if ne_type and str(ne_type).lower() != 'nan':
                    change_scope_lines.append(f"{count} * {ne_type}")
            change_scope = "\n".join(change_scope_lines)
        elif 'region_col' in cfg and cfg['region_col'] in group_data.columns:
            region_val = str(group_data.iloc[0].get(cfg['region_col'], ''))
            if region_val.lower() != 'nan' and region_val:
                change_scope = f"{region_val} region"
            else:
                change_scope = "All region"
        else:
            change_scope = "All region"
        global_replace(doc, '{{CHANGE_SCOPE}}', change_scope)
        
        # Handle CHANGE_RING_121 placeholder for New Integration PAG to AG scope
        if 'ring_121_col' in cfg and cfg['ring_121_col'] in group_data.columns:
            ring_121_val = str(group_data.iloc[0].get(cfg['ring_121_col'], ''))
            if ring_121_val.lower() != 'nan' and ring_121_val:
                global_replace(doc, '{{CHANGE_RING_121}}', ring_121_val)
            else:
                global_replace(doc, '{{CHANGE_RING_121}}', group_name)
        else:
            # If no ring_121_col, use the group name (Ring value)
            global_replace(doc, '{{CHANGE_RING_121}}', str(group_name))
        
        # Insert topology image for this specific ring (if provided)
        ring_image_path = topology_images.get(str(group_name))
        if ring_image_path and cfg.get('has_topology_image', False) and os.path.exists(ring_image_path):
            try:
                image_inserted = False
                # Find placeholder text in document and replace with image
                for para in doc.paragraphs:
                    if '{{TOPOLOGY_IMAGE}}' in para.text:
                        para.clear()
                        run = para.add_run()
                        run.add_picture(ring_image_path, width=Inches(6))
                        log(f"   [IMAGE] Topology inserted for {group_name}")
                        image_inserted = True
                        break
                
                if not image_inserted:
                    # If no placeholder found in paragraphs, try to find in tables
                    for table in doc.tables:
                        if image_inserted:
                            break
                        for row in table.rows:
                            if image_inserted:
                                break
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if '{{TOPOLOGY_IMAGE}}' in para.text:
                                        para.clear()
                                        run = para.add_run()
                                        run.add_picture(ring_image_path, width=Inches(5.5))
                                        log(f"   [IMAGE] Topology inserted in table for {group_name}")
                                        image_inserted = True
                                        break
            except Exception as e:
                log(f"   [WARN] Gagal insert image untuk {group_name}: {str(e)}")
        elif cfg.get('has_topology_image', False):
            # Remove placeholder if no image provided for this ring
            global_replace(doc, '{{TOPOLOGY_IMAGE}}', '[Gambar topologi tidak tersedia]')
            log(f"   [INFO] No image for {group_name}")
        
        target_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                header_row_text = " ".join([cell.text for cell in table.rows[0].cells])
                if cfg['table_keyword'] in header_row_text:
                    target_table = table
                    break
        
        if target_table:
            target_table.style = 'Table Grid'
            
            for i in range(len(target_table.rows)-1, 0, -1):
                row = target_table.rows[i]
                tbl = row._element.getparent()
                tbl.remove(row._element)
            
            cols_to_fill = cfg['columns_mapping']
            
            for idx, row in group_data.iterrows():
                new_row = target_table.add_row()
                row_cells = new_row.cells
                
                start_data_idx = 0
                if len(row_cells) > len(cols_to_fill):
                    row_cells[0].text = str(list(group_data.index).index(idx) + 1)
                    start_data_idx = 1
                
                for i, col_name in enumerate(cols_to_fill):
                    if start_data_idx + i < len(row_cells):
                        val = str(row.get(col_name, ''))
                        if val.lower() == 'nan': val = ''
                        row_cells[start_data_idx + i].text = val
                
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
            
            reset_table_width(target_table)
        
        clean_name = sanitize_filename(doc_title)
        output_path = os.path.join(output_folder, f"{clean_name}.docx")
        doc.save(output_path)
        
        # Auto-compress if file > 1 MB
        compressed, orig_size, new_size = compress_docx(output_path, max_size_mb=1)
        if compressed:
            reduction = ((orig_size - new_size) / orig_size) * 100
            log(f"   [COMPRESS] {os.path.basename(output_path)}: {orig_size/1024/1024:.2f}MB → {new_size/1024/1024:.2f}MB ({reduction:.1f}% smaller)")
        
        generated_files.append(output_path)
    
    zip_filename = f"Hasil_{scope_key}_{timestamp}.zip"
    zip_path = os.path.join(app.config['OUTPUT_FOLDER'], zip_filename)
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file in generated_files:
            zipf.write(file, os.path.basename(file))
    
    log(f"[ZIP] Created: {zip_filename}")
    log(f"[DONE] Total {len(generated_files)} dokumen berhasil")
    
    return zip_path, f"Berhasil generate {len(generated_files)} dokumen", logs, len(generated_files)

# ==========================================
#           API ROUTES
# ==========================================

@app.route('/')
def index():
    config = load_config()
    scopes = config.get("scopes", {})
    return render_template('index.html', scopes=scopes, page='generate')

@app.route('/templates')
def templates_page():
    config = load_config()
    scopes = config.get("scopes", {})
    # List uploaded templates
    templates = []
    if os.path.exists(app.config['TEMPLATES_FOLDER']):
        for f in os.listdir(app.config['TEMPLATES_FOLDER']):
            if f.endswith('.docx'):
                templates.append({
                    'name': f,
                    'path': os.path.join(app.config['TEMPLATES_FOLDER'], f),
                    'size': os.path.getsize(os.path.join(app.config['TEMPLATES_FOLDER'], f))
                })
    return render_template('index.html', scopes=scopes, templates=templates, page='templates')

@app.route('/config')
def config_page():
    config = load_config()
    scopes = config.get("scopes", {})
    return render_template('index.html', scopes=scopes, page='config')

@app.route('/history')
def history_page():
    config = load_config()
    scopes = config.get("scopes", {})
    history = load_history()
    return render_template('index.html', scopes=scopes, history=history.get("history", []), page='history')

@app.route('/excel-editor')
def excel_editor_page():
    config = load_config()
    scopes = config.get("scopes", {})
    return render_template('index.html', scopes=scopes, page='excel_editor')

# API: Upload and Generate
@app.route('/api/generate', methods=['POST'])
def api_generate():
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded', 'logs': []})
    
    file = request.files['file']
    scope = request.form.get('scope', '')
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected', 'logs': []})
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'success': False, 'error': 'File harus Excel (.xlsx atau .xls)', 'logs': []})
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    
    # Handle multiple topology images (one per ring)
    # Images are sent with keys like 'topology_Jabo Ring 35', 'topology_Jabo Ring 44', etc.
    topology_images = {}
    saved_image_paths = []
    
    for key in request.files:
        if key.startswith('topology_'):
            ring_name = key.replace('topology_', '')
            topo_file = request.files[key]
            if topo_file and topo_file.filename:
                # Create safe filename with ring name
                safe_ring_name = re.sub(r'[^\w\s-]', '', ring_name).replace(' ', '_')
                topo_filename = secure_filename(topo_file.filename)
                topo_path = os.path.join(app.config['UPLOAD_FOLDER'], f"topo_{safe_ring_name}_{topo_filename}")
                topo_file.save(topo_path)
                topology_images[ring_name] = topo_path
                saved_image_paths.append(topo_path)
    
    zip_path, message, logs, doc_count = generate_documents(filepath, scope, topology_images)
    
    # Cleanup uploaded files
    os.remove(filepath)
    for img_path in saved_image_paths:
        if os.path.exists(img_path):
            os.remove(img_path)
    
    if zip_path:
        add_to_history(scope, filename, doc_count, os.path.basename(zip_path))
        return jsonify({
            'success': True, 
            'message': message,
            'download_url': f'/download/{os.path.basename(zip_path)}',
            'logs': logs,
            'doc_count': doc_count
        })
    else:
        add_to_history(scope, filename, 0, "", "error")
        return jsonify({'success': False, 'error': message, 'logs': logs})

# API: Upload Template
@app.route('/api/templates/upload', methods=['POST'])
def api_upload_template():
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'})
    
    file = request.files['file']
    if not file.filename.endswith('.docx'):
        return jsonify({'success': False, 'error': 'File harus Word (.docx)'})
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['TEMPLATES_FOLDER'], filename)
    file.save(filepath)
    
    return jsonify({'success': True, 'message': f'Template "{filename}" berhasil diupload', 'path': filepath})

# API: Delete Template
@app.route('/api/templates/delete/<filename>', methods=['DELETE'])
def api_delete_template(filename):
    filepath = os.path.join(app.config['TEMPLATES_FOLDER'], secure_filename(filename))
    if os.path.exists(filepath):
        os.remove(filepath)
        return jsonify({'success': True, 'message': f'Template "{filename}" dihapus'})
    return jsonify({'success': False, 'error': 'Template tidak ditemukan'})

# API: Get Scopes
@app.route('/api/scopes')
def api_get_scopes():
    config = load_config()
    return jsonify(config.get("scopes", {}))

# API: Save Scope
@app.route('/api/scopes/<scope_key>', methods=['POST'])
def api_save_scope(scope_key):
    data = request.json
    config = load_config()
    config["scopes"][scope_key] = data
    save_config(config)
    return jsonify({'success': True, 'message': f'Scope "{scope_key}" saved'})

# API: Delete Scope
@app.route('/api/scopes/<scope_key>', methods=['DELETE'])
def api_delete_scope(scope_key):
    config = load_config()
    if scope_key in config.get("scopes", {}):
        del config["scopes"][scope_key]
        save_config(config)
        return jsonify({'success': True, 'message': f'Scope "{scope_key}" deleted'})
    return jsonify({'success': False, 'error': 'Scope tidak ditemukan'})

# ==========================================
#           EXCEL EDITOR APIs
# ==========================================

@app.route('/api/excel/parse', methods=['POST'])
def api_excel_parse():
    """Parse Excel file and return data as JSON for editing"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'})
    
    file = request.files['file']
    sheet_name = request.form.get('sheet', None)  # Optional sheet name
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'})
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'success': False, 'error': 'File harus Excel (.xlsx atau .xls)'})
    
    # Save file temporarily
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"edit_{filename}")
    file.save(filepath)
    
    try:
        # Get sheet names first
        xl = pd.ExcelFile(filepath)
        sheet_names = xl.sheet_names
        
        # Read the specified sheet or first sheet
        target_sheet = sheet_name if sheet_name and sheet_name in sheet_names else sheet_names[0]
        df = pd.read_excel(xl, sheet_name=target_sheet, dtype=str)
        df = df.fillna('')  # Replace NaN with empty string
        
        # Close ExcelFile to release file handle (important for Windows!)
        xl.close()
        
        # Convert to JSON-friendly format
        headers = df.columns.tolist()
        rows = df.values.tolist()
        
        # Delete temp file (with retry for Windows file locking)
        try:
            os.remove(filepath)
        except PermissionError:
            pass  # File may still be locked, will be cleaned up later
        
        return jsonify({
            'success': True,
            'filename': file.filename,
            'sheet_names': sheet_names,
            'current_sheet': target_sheet,
            'headers': headers,
            'rows': rows,
            'row_count': len(rows),
            'col_count': len(headers)
        })
        
    except Exception as e:
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
        except:
            pass  # Ignore cleanup errors
        return jsonify({'success': False, 'error': f"Error membaca Excel: {str(e)}"})

@app.route('/api/excel/save', methods=['POST'])
def api_excel_save():
    """Save edited data back to Excel file"""
    try:
        data = request.json
        headers = data.get('headers', [])
        rows = data.get('rows', [])
        filename = data.get('filename', 'edited_data.xlsx')
        
        if not headers or not rows:
            return jsonify({'success': False, 'error': 'No data to save'})
        
        # Create DataFrame
        df = pd.DataFrame(rows, columns=headers)
        
        # Save to Excel
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = secure_filename(filename.replace('.xlsx', '').replace('.xls', ''))
        output_filename = f"{safe_filename}_edited_{timestamp}.xlsx"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        df.to_excel(output_path, index=False, sheet_name='Sheet1')
        
        return jsonify({
            'success': True,
            'message': f'Excel saved: {output_filename}',
            'download_url': f'/download/{output_filename}',
            'filename': output_filename
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f"Error saving Excel: {str(e)}"})

# ==========================================
#       AUTO-FILL APIs (NE Report based)
# ==========================================

@app.route('/api/excel/parse-ne-report', methods=['POST'])
def api_excel_parse_ne_report():
    """Parse NE Report from Sheet2 to get NE Name -> Subnet mapping for Ring auto-fill"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'})
    
    file = request.files['file']
    sheet_name = request.form.get('sheet', 'Sheet2')
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'})
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'success': False, 'error': 'File harus Excel (.xlsx atau .xls)'})
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"ne_report_{filename}")
    file.save(filepath)
    
    try:
        xl = pd.ExcelFile(filepath)
        
        if sheet_name not in xl.sheet_names:
            xl.close()
            os.remove(filepath)
            return jsonify({'success': False, 'error': f"Sheet '{sheet_name}' tidak ditemukan. Available: {', '.join(xl.sheet_names)}"})
        
        df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)
        df.columns = df.columns.str.strip()
        xl.close()
        
        # Find NE Name column (flexible matching)
        ne_name_actual = None
        subnet_actual = None
        
        for col in df.columns:
            col_lower = col.lower()
            if 'ne name' in col_lower or 'nename' in col_lower or col_lower == 'ne name':
                ne_name_actual = col
            if 'subnet' in col_lower:
                subnet_actual = col
        
        if not ne_name_actual:
            os.remove(filepath)
            return jsonify({'success': False, 'error': f"Kolom 'NE Name' tidak ditemukan. Kolom tersedia: {', '.join(df.columns.tolist())}"})
        if not subnet_actual:
            os.remove(filepath)
            return jsonify({'success': False, 'error': f"Kolom 'Subnet' tidak ditemukan. Kolom tersedia: {', '.join(df.columns.tolist())}"})
        
        # Build NE report data
        ne_report = []
        for _, row in df.iterrows():
            ne_name = str(row[ne_name_actual]).strip() if pd.notna(row[ne_name_actual]) else ''
            subnet = str(row[subnet_actual]).strip() if pd.notna(row[subnet_actual]) else ''
            
            if ne_name and ne_name.lower() != 'nan':
                ne_report.append({
                    'ne_name': ne_name,
                    'subnet': subnet
                })
        
        os.remove(filepath)
        
        return jsonify({
            'success': True,
            'ne_report': ne_report,
            'record_count': len(ne_report),
            'columns_found': {
                'ne_name_col': ne_name_actual,
                'subnet_col': subnet_actual
            }
        })
        
    except Exception as e:
        if os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'success': False, 'error': f"Error parsing NE Report: {str(e)}"})

@app.route('/api/excel/autofill-ring', methods=['POST'])
def api_excel_autofill_ring():
    """Auto-fill Ring column based on Hostname matching NE Report data"""
    data = request.get_json()
    
    headers = data.get('headers', [])
    rows = data.get('rows', [])
    ne_report = data.get('ne_report', [])
    hostname_col_index = data.get('hostname_col_index', -1)
    ring_col_index = data.get('ring_col_index', -1)
    
    if hostname_col_index < 0 or ring_col_index < 0:
        return jsonify({'success': False, 'error': 'Invalid column indices'})
    
    # Build lookup dictionary: NE Name (hostname) -> Subnet (ring)
    ne_lookup = {}
    for record in ne_report:
        ne_name = str(record.get('ne_name', '')).strip()
        subnet = str(record.get('subnet', '')).strip()
        if ne_name and ne_name.lower() != 'nan':
            ne_lookup[ne_name.lower()] = subnet
    
    filled_count = 0
    not_found = []
    
    for row in rows:
        hostname = str(row[hostname_col_index]).strip() if hostname_col_index < len(row) else ''
        
        if not hostname or hostname.lower() == 'nan':
            continue
        
        current_ring = str(row[ring_col_index]).strip() if ring_col_index < len(row) else ''
        if current_ring and current_ring.lower() != 'nan':
            continue  # Skip if already has value
        
        ring_value = ne_lookup.get(hostname.lower())
        
        if ring_value:
            # Ensure row has enough columns
            while len(row) <= ring_col_index:
                row.append('')
            row[ring_col_index] = ring_value
            filled_count += 1
        else:
            if hostname not in not_found:
                not_found.append(hostname)
    
    return jsonify({
        'success': True,
        'rows': rows,
        'message': f'Auto-fill Ring selesai. {filled_count} baris terisi.',
        'filled_count': filled_count,
        'not_found': not_found[:20]  # Limit to first 20
    })

@app.route('/api/excel/parse-hostname-report', methods=['POST'])
def api_excel_parse_hostname_report():
    """Parse NE Report from Sheet2 to get Site ID -> NE Name (Hostname) mapping"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'})
    
    file = request.files['file']
    sheet_name = request.form.get('sheet', 'Sheet2')
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'})
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'success': False, 'error': 'File harus Excel (.xlsx atau .xls)'})
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"hostname_report_{filename}")
    file.save(filepath)
    
    try:
        xl = pd.ExcelFile(filepath)
        
        if sheet_name not in xl.sheet_names:
            xl.close()
            os.remove(filepath)
            return jsonify({'success': False, 'error': f"Sheet '{sheet_name}' tidak ditemukan. Available: {', '.join(xl.sheet_names)}"})
        
        df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)
        df.columns = df.columns.str.strip()
        xl.close()
        
        # Find Site ID and NE Name columns (flexible matching)
        site_id_actual = None
        ne_name_actual = None
        
        for col in df.columns:
            col_lower = col.lower()
            if 'site id' in col_lower or 'site_id' in col_lower or 'siteid' in col_lower or col_lower == 'site id':
                site_id_actual = col
            if 'ne name' in col_lower or 'nename' in col_lower or col_lower == 'ne name' or 'hostname' in col_lower:
                ne_name_actual = col
        
        if not site_id_actual:
            os.remove(filepath)
            return jsonify({'success': False, 'error': f"Kolom 'Site ID' tidak ditemukan. Kolom tersedia: {', '.join(df.columns.tolist())}"})
        if not ne_name_actual:
            os.remove(filepath)
            return jsonify({'success': False, 'error': f"Kolom 'NE Name' tidak ditemukan. Kolom tersedia: {', '.join(df.columns.tolist())}"})
        
        # Build hostname report data
        hostname_report = []
        for _, row in df.iterrows():
            site_id = str(row[site_id_actual]).strip() if pd.notna(row[site_id_actual]) else ''
            ne_name = str(row[ne_name_actual]).strip() if pd.notna(row[ne_name_actual]) else ''
            
            if site_id and site_id.lower() != 'nan' and ne_name and ne_name.lower() != 'nan':
                hostname_report.append({
                    'site_id': site_id,
                    'ne_name': ne_name
                })
        
        os.remove(filepath)
        
        return jsonify({
            'success': True,
            'hostname_report': hostname_report,
            'record_count': len(hostname_report),
            'columns_found': {
                'site_id_col': site_id_actual,
                'ne_name_col': ne_name_actual
            }
        })
        
    except Exception as e:
        if os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'success': False, 'error': f"Error parsing Hostname Report: {str(e)}"})

@app.route('/api/excel/autofill-hostname', methods=['POST'])
def api_excel_autofill_hostname():
    """Auto-fill Hostname column based on Site ID matching NE Report data"""
    data = request.get_json()
    
    headers = data.get('headers', [])
    rows = data.get('rows', [])
    hostname_report = data.get('hostname_report', [])
    site_id_col_index = data.get('site_id_col_index', -1)
    hostname_col_index = data.get('hostname_col_index', -1)
    
    if site_id_col_index < 0 or hostname_col_index < 0:
        return jsonify({'success': False, 'error': 'Invalid column indices'})
    
    # Build lookup dictionary: Site ID -> NE Name (Hostname)
    hostname_lookup = {}
    for record in hostname_report:
        site_id = str(record.get('site_id', '')).strip()
        ne_name = str(record.get('ne_name', '')).strip()
        if site_id and site_id.lower() != 'nan':
            hostname_lookup[site_id.lower()] = ne_name
    
    filled_count = 0
    not_found = []
    
    for row in rows:
        site_id = str(row[site_id_col_index]).strip() if site_id_col_index < len(row) else ''
        
        if not site_id or site_id.lower() == 'nan':
            continue
        
        current_hostname = str(row[hostname_col_index]).strip() if hostname_col_index < len(row) else ''
        if current_hostname and current_hostname.lower() != 'nan':
            continue  # Skip if already has value
        
        hostname_value = hostname_lookup.get(site_id.lower())
        
        if hostname_value:
            # Ensure row has enough columns
            while len(row) <= hostname_col_index:
                row.append('')
            row[hostname_col_index] = hostname_value
            filled_count += 1
        else:
            if site_id not in not_found:
                not_found.append(site_id)
    
    return jsonify({
        'success': True,
        'rows': rows,
        'message': f'Auto-fill Hostname selesai. {filled_count} baris terisi.',
        'filled_count': filled_count,
        'not_found': not_found[:20]  # Limit to first 20
    })



# API: Parse Rings from Excel
@app.route('/api/parse-rings', methods=['POST'])
def api_parse_rings():
    """Parse Excel file and extract unique ring values for topology image upload"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'})
    
    file = request.files['file']
    scope = request.form.get('scope', '')
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'})
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'success': False, 'error': 'File harus Excel (.xlsx atau .xls)'})
    
    config = load_config()
    cfg = config.get("scopes", {}).get(scope)
    
    if not cfg:
        return jsonify({'success': False, 'error': f"Scope '{scope}' tidak ditemukan"})
    
    # Check if scope supports topology image
    if not cfg.get('has_topology_image', False):
        return jsonify({'success': False, 'error': 'Scope ini tidak memerlukan gambar topologi'})
    
    # Save file temporarily
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{filename}")
    file.save(filepath)
    
    try:
        df = pd.read_excel(filepath, sheet_name=cfg['excel_sheet'], dtype=str)
        df.columns = df.columns.str.strip()
        
        ring_col = cfg.get('ring_col', 'Ring')
        if ring_col not in df.columns:
            os.remove(filepath)
            return jsonify({'success': False, 'error': f"Kolom '{ring_col}' tidak ditemukan di Excel"})
        
        # Extract unique rings
        df[ring_col] = df[ring_col].str.strip()
        unique_rings = df[ring_col].dropna().unique().tolist()
        unique_rings = [r for r in unique_rings if r and str(r).lower() != 'nan']
        unique_rings.sort()
        
        os.remove(filepath)
        
        return jsonify({
            'success': True,
            'rings': unique_rings,
            'total': len(unique_rings)
        })
        
    except Exception as e:
        if os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'success': False, 'error': f"Error membaca Excel: {str(e)}"})

# API: Get History
@app.route('/api/history')
def api_get_history():
    return jsonify(load_history())

# API: Clear History
@app.route('/api/history/clear', methods=['POST'])
def api_clear_history():
    save_history({"history": []})
    return jsonify({'success': True, 'message': 'History cleared'})

# ==========================================
#           XLWO FINDER
# ==========================================

XLWO_CONFIG_FILE = 'xlwo_config.json'
XLWO_DATA_FOLDER = 'xlwo_data'

# Ensure XLWO data folder exists
os.makedirs(XLWO_DATA_FOLDER, exist_ok=True)

def load_xlwo_config():
    """Load XLWO database configuration"""
    if os.path.exists(XLWO_CONFIG_FILE):
        with open(XLWO_CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"database_file": None, "uploaded_at": None, "record_count": 0, "columns": []}

def save_xlwo_config(config):
    """Save XLWO database configuration"""
    with open(XLWO_CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2, ensure_ascii=False)

@app.route('/xlwo-finder')
def xlwo_finder_page():
    """Render XLWO Finder page"""
    config = load_config()
    scopes = config.get("scopes", {})
    xlwo_config = load_xlwo_config()
    return render_template('index.html', scopes=scopes, xlwo_config=xlwo_config, page='xlwo_finder')

@app.route('/api/xlwo/config')
def api_xlwo_config():
    """Get current XLWO database configuration"""
    config = load_xlwo_config()
    # Add last modified time if file exists
    if config.get('source_path') and os.path.exists(config['source_path']):
        mtime = os.path.getmtime(config['source_path'])
        config['last_modified'] = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M:%S')
    elif config.get('database_file') and os.path.exists(config['database_file']):
        mtime = os.path.getmtime(config['database_file'])
        config['last_modified'] = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M:%S')
    return jsonify(config)

@app.route('/api/xlwo/set-source', methods=['POST'])
def api_xlwo_set_source():
    """Set source Excel file path for direct file linking"""
    data = request.get_json()
    source_path = data.get('path', '').strip()
    
    if not source_path:
        return jsonify({'success': False, 'error': 'Path is required'})
    
    # Normalize path
    source_path = os.path.abspath(source_path)
    
    if not os.path.exists(source_path):
        return jsonify({'success': False, 'error': f'File not found: {source_path}'})
    
    if not source_path.lower().endswith(('.xlsx', '.xls')):
        return jsonify({'success': False, 'error': 'File must be Excel (.xlsx or .xls)'})
    
    try:
        # Read Excel to get column info and record count
        df = pd.read_excel(source_path)
        df.columns = df.columns.str.strip()
        columns = df.columns.tolist()
        record_count = len(df)
        
        # Get last modified time
        mtime = os.path.getmtime(source_path)
        last_modified = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M:%S')
        
        # Save configuration with source_path (not copying file)
        xlwo_config = {
            "source_path": source_path,  # Direct link to source file
            "database_file": source_path,  # For backward compatibility
            "original_filename": os.path.basename(source_path),
            "uploaded_at": datetime.now().isoformat(),
            "last_modified": last_modified,
            "record_count": record_count,
            "columns": columns,
            "mode": "direct"  # Mark as direct file link mode
        }
        save_xlwo_config(xlwo_config)
        
        return jsonify({
            'success': True,
            'message': f'Source linked: {record_count} records',
            'config': xlwo_config
        })
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error reading file: {str(e)}'})

@app.route('/api/xlwo/set-gsheet', methods=['POST'])
def api_xlwo_set_gsheet():
    """Set Google Sheets URL for cloud sync"""
    import requests
    import re
    
    data = request.get_json()
    gsheet_url = data.get('url', '').strip()
    
    if not gsheet_url:
        return jsonify({'success': False, 'error': 'Google Sheets URL is required'})
    
    try:
        # Extract spreadsheet ID from various URL formats
        # Format 1: https://docs.google.com/spreadsheets/d/SPREADSHEET_ID/edit
        # Format 2: https://docs.google.com/spreadsheets/d/SPREADSHEET_ID/pub?output=xlsx
        match = re.search(r'/spreadsheets/d/([a-zA-Z0-9-_]+)', gsheet_url)
        if not match:
            return jsonify({'success': False, 'error': 'Invalid Google Sheets URL. Use the sharing link from Google Sheets.'})
        
        spreadsheet_id = match.group(1)
        
        # Convert to export URL (xlsx format)
        export_url = f'https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=xlsx'
        
        # Download the spreadsheet
        response = requests.get(export_url, timeout=30)
        if response.status_code != 200:
            return jsonify({'success': False, 'error': 'Failed to download. Make sure the spreadsheet is published to web or shared with "Anyone with link".'})
        
        # Save to local cache file
        cache_file = os.path.join(XLWO_DATA_FOLDER, 'gsheet_cache.xlsx')
        with open(cache_file, 'wb') as f:
            f.write(response.content)
        
        # Read and validate the data
        df = pd.read_excel(cache_file)
        df.columns = df.columns.str.strip()
        columns = df.columns.tolist()
        record_count = len(df)
        
        # Save configuration
        xlwo_config = {
            "gsheet_url": gsheet_url,
            "gsheet_id": spreadsheet_id,
            "database_file": cache_file,
            "original_filename": f"Google Sheet ({spreadsheet_id[:8]}...)",
            "uploaded_at": datetime.now().isoformat(),
            "record_count": record_count,
            "columns": columns,
            "mode": "gsheet"  # Mark as Google Sheets mode
        }
        save_xlwo_config(xlwo_config)
        
        return jsonify({
            'success': True,
            'message': f'Google Sheet linked: {record_count} records',
            'config': xlwo_config
        })
    except requests.exceptions.RequestException as e:
        return jsonify({'success': False, 'error': f'Network error: {str(e)}'})
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error: {str(e)}'})

@app.route('/api/xlwo/upload', methods=['POST'])
def api_xlwo_upload():
    """Upload Excel file as XLWO database"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'})
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'success': False, 'error': 'File must be Excel (.xlsx or .xls)'})
    
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(XLWO_DATA_FOLDER, 'database.xlsx')
        file.save(filepath)
        
        # Read Excel to get column info and record count
        df = pd.read_excel(filepath)
        # Strip whitespace from column names
        df.columns = df.columns.str.strip()
        columns = df.columns.tolist()
        record_count = len(df)
        
        # Save configuration
        xlwo_config = {
            "database_file": filepath,
            "original_filename": filename,
            "uploaded_at": datetime.now().isoformat(),
            "record_count": record_count,
            "columns": columns
        }
        save_xlwo_config(xlwo_config)
        
        return jsonify({
            'success': True,
            'message': f'Database uploaded: {record_count} records',
            'config': xlwo_config
        })
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error uploading database: {str(e)}'})

@app.route('/api/xlwo/search')
def api_xlwo_search():
    """Search XLWO database with filters"""
    xlwo_config = load_xlwo_config()
    
    if not xlwo_config.get('database_file') or not os.path.exists(xlwo_config.get('database_file', '')):
        return jsonify({'success': False, 'error': 'No database uploaded', 'data': []})
    
    try:
        # Get search parameters
        query = request.args.get('q', '').strip().lower()
        status = request.args.get('status', 'all')  # all, active, expired
        field = request.args.get('field', 'all')  # specific column or 'all'
        
        # Read database
        df_original = pd.read_excel(xlwo_config['database_file'])
        # Strip whitespace from column names
        df_original.columns = df_original.columns.str.strip()
        df = df_original.copy()
        
        # Find expiry column - check various possible column names
        expiry_columns = ['Exp Date', 'Exp Day', 'Masa Berlaku', 'Expiry Date', 'Valid Until', 'Tanggal Expired', 'Expired', 'ExpDate', 'ExpDay']
        expiry_col = None
        for col in expiry_columns:
            if col in df.columns:
                expiry_col = col
                break
        
        # Calculate stats from original data
        stats = {'total': len(df_original), 'active': 0, 'expired': 0}
        if expiry_col:
            today = datetime.now().date()
            df_original[expiry_col] = pd.to_datetime(df_original[expiry_col], errors='coerce')
            valid_dates = df_original[expiry_col].notna()
            stats['active'] = int((df_original.loc[valid_dates, expiry_col].dt.date >= today).sum())
            stats['expired'] = int((df_original.loc[valid_dates, expiry_col].dt.date < today).sum())
        
        # Convert all columns to string for searching
        df_str = df.astype(str)
        
        # Search filter
        if query:
            if field == 'all':
                # Search across all columns
                mask = df_str.apply(lambda row: row.str.lower().str.contains(query, na=False).any(), axis=1)
            else:
                # Search specific column
                if field in df.columns:
                    mask = df_str[field].str.lower().str.contains(query, na=False)
                else:
                    mask = pd.Series([True] * len(df))
            df = df[mask]
        
        # Status filter
        if status != 'all' and expiry_col:
            today = datetime.now().date()
            df[expiry_col] = pd.to_datetime(df[expiry_col], errors='coerce')
            
            if status == 'expired':
                df = df[df[expiry_col].dt.date < today]
            elif status == 'active':
                df = df[df[expiry_col].dt.date >= today]
        
        # Convert to records, handling datetime
        total_results = len(df)
        df = df.fillna('')
        for col in df.columns:
            if df[col].dtype == 'datetime64[ns]':
                df[col] = df[col].dt.strftime('%Y-%m-%d').replace('NaT', '')
        
        records = df.head(100).to_dict('records')  # Limit to 100 records for performance
        
        return jsonify({
            'success': True,
            'data': records,
            'total': total_results,
            'columns': list(df.columns),  # Use cleaned column names from dataframe
            'expiry_column': expiry_col,
            'stats': stats
        })
    except Exception as e:
        return jsonify({'success': False, 'error': f'Search error: {str(e)}', 'data': []})


# ==========================================
#           AUTO-FILL C/AG HOSTNAME
# ==========================================

@app.route('/api/excel/autofill-cag', methods=['POST'])
def api_excel_autofill_cag():
    """Auto-fill C/AG Hostname based on PAG Hostname using topo sheet.
    
    Expects JSON with:
    - headers: list of column headers
    - rows: list of row data
    - topo_data: list of topology chain strings (from topo sheet)
    - pag_col_index: index of PAG Hostname column
    - cag_col_index: index of C/AG Hostname column
    """
    try:
        data = request.json
        headers = data.get('headers', [])
        rows = data.get('rows', [])
        topo_data = data.get('topo_data', [])
        pag_col_index = data.get('pag_col_index')
        cag_col_index = data.get('cag_col_index')
        
        if pag_col_index is None or cag_col_index is None:
            return jsonify({'success': False, 'error': 'Column indices required'})
        
        if not topo_data:
            return jsonify({'success': False, 'error': 'Topo data required'})
        
        # Build lookup dictionary from topo data
        # Each topo row is a comma-separated chain of hostnames
        filled_count = 0
        
        for row in rows:
            pag_hostname = str(row[pag_col_index]).strip() if pag_col_index < len(row) else ''
            
            if not pag_hostname or pag_hostname.lower() == 'nan':
                continue
            
            # Skip if C/AG already filled
            current_cag = str(row[cag_col_index]).strip() if cag_col_index < len(row) else ''
            if current_cag and current_cag.lower() != 'nan':
                continue
            
            # Search for PAG Hostname in topo chains
            nearest_cag = find_nearest_cag(pag_hostname, topo_data)
            
            if nearest_cag:
                row[cag_col_index] = nearest_cag
                filled_count += 1
        
        return jsonify({
            'success': True,
            'message': f'Auto-filled {filled_count} C/AG Hostname(s)',
            'rows': rows,
            'filled_count': filled_count
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error: {str(e)}'})

def find_nearest_cag(pag_hostname, topo_data):
    """Find the nearest C/AG hostname from PAG hostname in topo chains.
    
    Args:
        pag_hostname: The PAG hostname to search for
        topo_data: List of comma-separated chains
        
    Returns:
        The nearest AG/C hostname or None
    """
    for chain in topo_data:
        if not chain or not isinstance(chain, str):
            continue
        
        # Split chain into hostnames
        hostnames = [h.strip() for h in chain.split(',') if h.strip()]
        
        # Find PAG position in chain
        pag_pos = None
        for i, hostname in enumerate(hostnames):
            if hostname == pag_hostname:
                pag_pos = i
                break
        
        if pag_pos is None:
            continue
        
        # Find nearest AG/C hostname (prefix AG- or C-)
        # Search left and right simultaneously
        left_distance = float('inf')
        right_distance = float('inf')
        left_cag = None
        right_cag = None
        
        # Search left
        for i in range(pag_pos - 1, -1, -1):
            hostname = hostnames[i]
            if hostname.startswith('AG-') or hostname.startswith('C-'):
                left_distance = pag_pos - i
                left_cag = hostname
                break
        
        # Search right
        for i in range(pag_pos + 1, len(hostnames)):
            hostname = hostnames[i]
            if hostname.startswith('AG-') or hostname.startswith('C-'):
                right_distance = i - pag_pos
                right_cag = hostname
                break
        
        # Return the nearest one
        if left_cag and right_cag:
            return left_cag if left_distance <= right_distance else right_cag
        elif left_cag:
            return left_cag
        elif right_cag:
            return right_cag
    
    return None

@app.route('/api/excel/parse-topo', methods=['POST'])
def api_excel_parse_topo():
    """Parse topo sheet from Excel file and return chain data"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'})
    
    file = request.files['file']
    sheet_name = request.form.get('sheet', 'topo')
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'})
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'success': False, 'error': 'File must be Excel'})
    
    # Save file temporarily
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"topo_{filename}")
    file.save(filepath)
    
    try:
        xl = pd.ExcelFile(filepath)
        
        # Check if topo sheet exists
        if sheet_name not in xl.sheet_names:
            xl.close()
            os.remove(filepath)
            return jsonify({
                'success': False, 
                'error': f'Sheet "{sheet_name}" not found. Available sheets: {xl.sheet_names}'
            })
        
        # Read topo sheet - it's a single column of comma-separated chains
        df = pd.read_excel(xl, sheet_name=sheet_name, header=None)
        xl.close()
        
        # Extract all non-empty values as chains
        topo_chains = []
        for col in df.columns:
            for val in df[col].dropna().astype(str):
                if val and val.lower() != 'nan' and ',' in val:
                    topo_chains.append(val.strip())
        
        os.remove(filepath)
        
        return jsonify({
            'success': True,
            'topo_data': topo_chains,
            'chain_count': len(topo_chains)
        })
        
    except Exception as e:
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
        except:
            pass
        return jsonify({'success': False, 'error': f'Error reading topo sheet: {str(e)}'})

@app.route('/api/xlwo/clear', methods=['POST'])
def api_xlwo_clear():
    """Clear XLWO database"""
    try:
        xlwo_config = load_xlwo_config()
        if xlwo_config.get('database_file') and os.path.exists(xlwo_config['database_file']):
            os.remove(xlwo_config['database_file'])
        
        save_xlwo_config({"database_file": None, "uploaded_at": None, "record_count": 0, "columns": []})
        return jsonify({'success': True, 'message': 'Database cleared'})
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error clearing database: {str(e)}'})

# Download
@app.route('/download/<filename>')
def download_file(filename):
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True, download_name=filename)
    return "File not found", 404

if __name__ == '__main__':
    print("=" * 50)
    print("Document Automation v2.0")
    print("Open browser: http://localhost:5020")
    print("=" * 50)
    app.run(debug=True, port=5020)
